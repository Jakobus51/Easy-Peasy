import numpy as np
import pandas as pd
import bs4
import time
import csv
import requests
import json
import threading
import sys
import os

from docx2pdf import convert
from pathlib import Path
from os import path
import shutil
import docx


from urllib.request import urlopen
from urllib.error import HTTPError
from urllib.error import URLError
from bs4 import BeautifulSoup as soup

#gets the URL and trows error if there is something wrong
def GetHTML(URL):
	#Try 5 times to get server adress otherwise break
	for attempt in range(5):
		try:
			html = urlopen(URL)
		except HTTPError as e:
			print("{} , attempt {}".format(e,attempt))
			time.sleep(6)
			continue

		except URLError as e:
			print("The server could not be found, attempt {}".format(attempt))
			time.sleep(6)
			continue

		else:
			return html
	else:
		sys.exit("Error with retrieving html")

#Returns the value of amount of of results on bol.com
def GetBolResult(product):
	
	product_string = product.replace(" ", "+")
	my_url = "https://www.bol.com/nl/nl/s/?searchtext={}".format(product_string)

	html =  GetHTML(my_url)
	page_soup = soup(html.read(), "html.parser")
	
	ratings = page_soup.findAll("div", {"class":"star-rating"})
	bolProducts = page_soup.findAll("div", {"class":"product-item__info hit-area"})

	ratings = [None, None, None]
	count = 0
	for bolproduct in bolProducts:
		#check if the sponsor header exist if it does not, the rating will be added to the list
		#We only want the first 3 ratings from a page
		try:
			test = bolproduct.find("div", {"class":"h-color-subtext u-pb--xs small_details"})
			if test.text != "Gesponsord":
				#Sometimes he finds sponsor headers while products are not sponsored, so that is the reason for this extra code
				if bolproduct.find("div", {"class":"star-rating"}).get('data-count') != None:
					ratings[count] = bolproduct.find("div", {"class":"star-rating"}).get('data-count')
				else:
					ratings[count] = 0				
				count+=1
				if count == 3:
					break
			
		except:
			if bolproduct.find("div", {"class":"star-rating"}).get('data-count') != None:
				ratings[count] = bolproduct.find("div", {"class":"star-rating"}).get('data-count')
			else:
				ratings[count] = 0	

			count+=1
			if count == 3:
				break
				
	
	resultWithText = page_soup.findAll("p", {"class":"total-results js_total_results"})
	if  resultWithText:	
		result = RemoveNonDigit(resultWithText)
	else:	
		result = None

	time.sleep(1)
	return result, ratings[0], ratings[1], ratings[2]


#Retreives the tokens to acces the BOL API
def GetAccessToken():

    url = "https://login.bol.com/token?grant_type=client_credentials"
    clientBase64 = "ZTZjMjc3ODEtZjdhNS00Zjg5LWE1OTUtYWNjYjJjNzgxYzM5OkFNQmlWa2sta24tWVNlaUFRQksxdVVnckpyNXBhbVdfUWw2MTMzZnh1d1A0eHJOOURhUW9NSkdQUFJmRDBjSkhpLVNXNi1XUEh3dnRkdG9Dd2dCVmJRYw=="

    payload = {}
    headers = {   
      'Authorization': "Basic " + clientBase64,  
        }

    response = requests.request("POST", url, headers=headers, data=payload)
    responseJson = json.loads(response.text)
    AccessToken = responseJson["access_token"]	
   
    return AccessToken



#Finds the 4 most searched synonims and also the searchvolume for 30 days for the original product and its synonims
def GetSearchResults(accesToken, product, time_frame, period):
	#Try 5 times to get server adress otherwise break
	for attempt in range(5):
		product_string = product.replace(" ", "+")
		url = "https://api.bol.com/retailer/insights/search-terms?search-term={}&period={}&number-of-periods={}&related-search-terms=true".format(product_string, time_frame, period)
		payload={}

		headers = {
		"Authorization": "Bearer " + accesToken,
		'Accept': 'application/vnd.retailer.v5+json'
		}

		response = requests.request("GET", url, headers=headers, data=payload)

		responsePython = json.loads(response.text)

		#Checks if we got back a good request, bad request contain a status element in their response
		try:
			responsePython["status"]
			print("Error Code: {}, Attempt: {}".format(responsePython["status"], attempt))
			time.sleep(6)
			continue

		except:
			return responsePython

	else:
		print("Error Code: {},\n Detail: {}".format(responsePython["status"], responsePython["detail"]))
		sys.exit("Error with talking to API error")


#Removes non digits
def RemoveNonDigit(result):
	numeric_filter = filter(str.isdigit, result[0].text)
	resultWithoutText = "".join(numeric_filter)
	return resultWithoutText


#Generate invoice by writing in the empty invoice and saving it as pdf
def generateInvoice(order):    
	doc = docx.Document("C:/Users/Jakob/Documents/EasyPeasy/Boekhouding/Empty invoice.docx")
	
	doc.paragraphs[10].text = "Naam:                  {}".format(order.name)
	doc.paragraphs[11].text = "Bedrijfsnaam:      {}".format(order.company)
	doc.paragraphs[12].text = "Datum:                {}".format(order.date)
	doc.paragraphs[13].text = "Factuur nummer: {}".format(order.invoiceNumber)
	doc.paragraphs[20].text = "Product onderzoek: {} pakket ({} producten)			  	€ {}".format(order.package, order.numberOfProducts,order.amount )
	doc.paragraphs[21].text = "BTW (21%)										€ {}".format(order.tax)    
	doc.paragraphs[22].text = "Totaal:										€ {}".format(order.total)
	doc.paragraphs[22].runs[0].bold = True
	doc.save("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{} {} factuur.docx".format( order.fileName, order.invoiceNumber, order.fileName))

	convert("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{} {} factuur.docx".format( order.fileName, order.invoiceNumber, order.fileName))
	shutil.copyfile("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{} {} factuur.pdf".format( order.fileName, order.invoiceNumber, order.fileName), "C:/Users/Jakob/Documents/EasyPeasy/Boekhouding/Facturen Uit/{} {} factuur.pdf".format(order.invoiceNumber, order.fileName))

	
#Write the results into an excel and colour the header
def writeExcel(order, products):
	
	writer = pd.ExcelWriter("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{} Resultaten ({}).xlsx".format(order.fileName, order.fileName, order.package), engine='xlsxwriter')

	products.to_excel(writer, sheet_name="Sheet1" , startrow=1, header=False, index= False)
	workbook  = writer.book
	worksheet = writer.sheets['Sheet1']
	
	header_format = workbook.add_format({
		'bold': True,			
		"align": "center",
		'fg_color': '{}'.format(order.colour),
		'border': 1
		})
	
	for col_num, value in enumerate(products.columns.values):
		worksheet.write(0, col_num , value, header_format)
	
	writer.save()

def setAndFillFolder(order):
		
	#Check if the file exist otherwise make a new one
	if (Path("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{}.csv".format(order.fileName, order.fileName)).exists()):
		print("file found: {}.csv".format(order.fileName))
		input("Want to check file? otherwise proceed")
	else:
		Path("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}".format(order.fileName, order.fileName)).mkdir(parents=True, exist_ok=True)
		#check if a correct list was retrieved from sample form
		if(order.package == "Sample"):
			df = pd.DataFrame(order.productList)
			df.to_csv("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{}.csv".format(order.fileName, order.fileName), index=False, header=False)
		
			print("+"*20)
			print(order.productList)
			print("+"*20)
						
			input("({}), is above list correct? Press enter to continue".format(order.fileName))
		
		#wait till the list is filled of the order
		else:
			df = pd.DataFrame([])
			df.to_csv("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{}.csv".format(order.fileName, order.fileName))
			print("Message of customer:")
			print(order.message)
			print()

			os.startfile("C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/{}.csv".format(order.fileName, order.fileName))
			input("({}), fill csv please. Press enter to continue".format(order.fileName))


def addInstruction(order):
	#Add the instruction pdf to the file
	if(order.package == "Sample" or order.package == "Complete"):
		shutil.copyfile("C:/Users/Jakob/Documents/EasyPeasy/Uitleg/Complete.pdf", "C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/Complete.pdf".format(order.fileName))

	if(order.package == "Advanced"):
		shutil.copyfile("C:/Users/Jakob/Documents/EasyPeasy/Uitleg/Advanced.pdf", "C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/Advanced.pdf".format(order.fileName))

	if(order.package == "Basic"):
		shutil.copyfile("C:/Users/Jakob/Documents/EasyPeasy/Uitleg/Basic.pdf", "C:/Users/Jakob/Documents/EasyPeasy/Resultaten/{}/Basic.pdf".format(order.fileName))




	